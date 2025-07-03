"""
Phân tích tài liệu Tiếng Việt - Vietnamese Document Parser

Module này cung cấp khả năng phân tích và trích xuất nội dung từ các loại tài liệu khác nhau,
được tối ưu hóa đặc biệt cho tiếng Việt.

Hỗ trợ các định dạng:
- PDF: Sử dụng PyMuPDF (ưu tiên) hoặc PyPDFLoader (fallback) + pdfplumber cho bảng
- DOCX: Trích xuất đoạn văn và bảng từ Word documents
- TXT/MD: Text files với nhiều encoding
- HTML: Thêm hỗ trợ cơ bản cho HTML files
- CSV: Thêm hỗ trợ cơ bản cho CSV files

Tính năng đặc biệt:
- Nhận diện và làm sạch text tiếng Việt
- Phân loại nội dung (bảng, tiêu đề, danh sách, v.v.)
- Trích xuất metadata chi tiết
- Error handling mạnh mẽ với fallback options
"""

import os
import logging
import json
from pathlib import Path
from typing import List, Dict, Any, Optional
import re
import csv
import io

# Loader cơ bản
try:
    from langchain_community.document_loaders import PyPDFLoader, TextLoader
except ImportError:
    from langchain.document_loaders import PyPDFLoader, TextLoader

try:
    from langchain_core.documents import Document
except ImportError:
    from langchain.schema import Document

# Thư viện phụ trợ (nếu có)
try:
    import fitz  # PyMuPDF để tách bố cục PDF

    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pdfplumber  # Tách bảng PDF

    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import pandas as pd  # Định dạng bảng đẹp hơn

    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    from docx import Document as DocxDocument  # DOCX support

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup  # HTML parsing

    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Trình phân tích tài liệu Tiếng Việt

    Class chính để xử lý các loại tài liệu khác nhau, được tối ưu hóa cho tiếng Việt.
    Tự động phát hiện và sử dụng các thư viện tốt nhất có sẵn.

    Attributes:
        supported_formats (set): Tập hợp các định dạng file được hỗ trợ
        capabilities (dict): Dictionary chứa thông tin các thư viện có sẵn
    """

    def __init__(self):
        """
        Khởi tạo DocumentParser

        Kiểm tra các thư viện có sẵn và cấu hình định dạng hỗ trợ.
        Log ra thông tin các tính năng có thể sử dụng.
        """
        # Định dạng file hỗ trợ cơ bản
        self.supported_formats = {".pdf", ".txt", ".md", ".csv"}

        # Thêm định dạng tùy theo thư viện có sẵn
        if DOCX_AVAILABLE:
            self.supported_formats.add(".docx")
        if BS4_AVAILABLE:
            self.supported_formats.update({".html", ".htm"})

        # Kiểm tra thư viện nào đã cài
        self.capabilities = self._check_capabilities()

        # Cache dictionary để tránh load nhiều lần
        self._vietnamese_words_cache = None
        self._vietnamese_abbreviations_cache = None

        # Log ra hỗ trợ
        logger.info(f"Định dạng hỗ trợ: {', '.join(sorted(self.supported_formats))}")
        self._log_capabilities()

    def _check_capabilities(self) -> Dict[str, bool]:
        """
        Kiểm tra các thư viện phụ trợ có sẵn

        Returns:
            Dict[str, bool]: Dictionary chứa tên thư viện và trạng thái có sẵn
                - pymupdf: Có PyMuPDF cho PDF nâng cao
                - pdfplumber: Có pdfplumber cho trích xuất bảng PDF
                - pandas: Có pandas cho xử lý dữ liệu
                - docx: Có python-docx cho file Word
                - bs4: Có BeautifulSoup cho HTML
        """
        return {
            "pymupdf": PYMUPDF_AVAILABLE,
            "pdfplumber": PDFPLUMBER_AVAILABLE,
            "pandas": PANDAS_AVAILABLE,
            "docx": DOCX_AVAILABLE,
            "bs4": BS4_AVAILABLE,
        }

    def _log_capabilities(self):
        """
        Log ra trạng thái các thư viện phụ trợ

        In ra thông tin về việc có thể sử dụng các tính năng nâng cao hay không.
        Giúp debug khi có vấn đề với dependencies.
        """
        if PYMUPDF_AVAILABLE:
            logger.info("Có PyMuPDF (phân tích PDF nâng cao)")
        else:
            logger.warning("Thiếu PyMuPDF - sẽ dùng PyPDFLoader cơ bản")

        if PDFPLUMBER_AVAILABLE:
            logger.info("Có pdfplumber (tách bảng PDF)")
        else:
            logger.warning("Thiếu pdfplumber - không tách được bảng PDF")

        if PANDAS_AVAILABLE:
            logger.info("Có pandas (xử lý bảng)")
        else:
            logger.warning("Thiếu pandas - hạn chế xử lý bảng")

        if DOCX_AVAILABLE:
            logger.info("Có python-docx (hỗ trợ DOCX)")
        else:
            logger.warning("Thiếu python-docx - không đọc được file Word")

        if BS4_AVAILABLE:
            logger.info("Có BeautifulSoup (hỗ trợ HTML)")
        else:
            logger.warning("Thiếu BeautifulSoup - không đọc được HTML")

    def _contains_vietnamese(self, text: str) -> bool:
        """
        Kiểm tra text có chứa ký tự tiếng Việt hay không

        Args:
            text (str): Đoạn text cần kiểm tra

        Returns:
            bool: True nếu có ký tự tiếng Việt, False nếu không

        Examples:
            >>> parser._contains_vietnamese("Xin chào")
            True
            >>> parser._contains_vietnamese("Hello world")
            False
        """
        vietnamese_chars = (
            r"[àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđĐ]"
        )
        return bool(re.search(vietnamese_chars, text, re.IGNORECASE))

    def _get_vietnamese_words(self) -> set:
        """
        Load từ điển tiếng Việt từ file JSON external (với cache)

        Returns:
            set: Set các từ tiếng Việt thường gặp, load từ vietnamese_words.json

        Note:
            - Cache kết quả để tránh load file nhiều lần
            - Ưu tiên load từ file JSON để dễ mở rộng
            - Fallback về hardcoded list nếu file không tồn tại
        """
        # Return cache nếu đã load
        if self._vietnamese_words_cache is not None:
            return self._vietnamese_words_cache

        try:
            # Tìm file vietnamese_words.json
            current_file = Path(__file__).resolve()
            aiservice_root = current_file.parents[2]
            json_file = aiservice_root / "data" / "vietnamese_words.json"

            if json_file.exists():
                with open(json_file, "r", encoding="utf-8") as f:
                    data = json.load(f)

                # Combine tất cả categories thành một set
                all_words = set()
                for category, words in data.items():
                    if category != "metadata" and isinstance(words, list):
                        all_words.update(words)

                logger.info(
                    f"Loaded {len(all_words)} Vietnamese words from {json_file.name}"
                )
                self._vietnamese_words_cache = all_words
                return all_words
            else:
                logger.warning(f"Vietnamese words file not found: {json_file}")

        except Exception as e:
            logger.error(f"Error loading Vietnamese words from JSON: {e}")

        # Fallback: hardcoded essential words
        logger.info("Using fallback Vietnamese words dictionary")
        fallback_words = {
            # Essential pronouns
            "tôi",
            "bạn",
            "anh",
            "chị",
            "em",
            "mình",
            "ta",
            "họ",
            "ai",
            "gì",
            # Essential verbs
            "là",
            "có",
            "được",
            "làm",
            "đi",
            "đến",
            "về",
            "xin",
            "chào",
            # Essential adjectives
            "tốt",
            "đẹp",
            "lớn",
            "nhỏ",
            "mới",
            "cũ",
            # Essential connectors
            "và",
            "với",
            "từ",
            "để",
            "trong",
            "ngoài",
            # Essential particles
            "ạ",
            "ơi",
            "nhé",
            "rồi",
            "đã",
            "chưa",
            # Essential professional words
            "lý",
            "viên",
            "công",
            "ty",
            "trợ",
            "giúp",
        }
        self._vietnamese_words_cache = fallback_words
        return fallback_words

    def _get_vietnamese_abbreviations(self) -> set:
        """
        Load abbreviations (viết tắt) từ file JSON (với cache)

        Returns:
            set: Set các từ viết tắt được coi là tiếng Việt trong context
        """
        # Return cache nếu đã load
        if self._vietnamese_abbreviations_cache is not None:
            return self._vietnamese_abbreviations_cache

        try:
            current_dir = Path(__file__).parent
            json_file = current_dir / "vietnamese_words.json"

            if json_file.exists():
                with open(json_file, "r", encoding="utf-8") as f:
                    data = json.load(f)

                abbreviations = data.get("abbreviations", [])
                abbreviations_set = set(abbreviations)
                self._vietnamese_abbreviations_cache = abbreviations_set
                return abbreviations_set

        except Exception as e:
            logger.error(f"Error loading Vietnamese abbreviations: {e}")

        # Fallback
        fallback_abbr = {"ai", "it", "hr", "ceo"}
        self._vietnamese_abbreviations_cache = fallback_abbr
        return fallback_abbr

    def _calculate_vietnamese_score(self, text: str) -> float:
        """
        Tính tỷ lệ tiếng Việt trong text (thuật toán thông minh)

        Args:
            text (str): Đoạn text cần tính toán

        Returns:
            float: Tỷ lệ từ 0.0 đến 1.0
                  (0.0 = không có tiếng Việt, 1.0 = toàn bộ là tiếng Việt)

        Examples:
            >>> parser._calculate_vietnamese_score("Xin chào các bạn")
            0.9  # 90% tiếng Việt (có từ vựng + dấu)
            >>> parser._calculate_vietnamese_score("Hello world")
            0.0  # Không có tiếng Việt

        Logic thông minh:
            - 70% trọng số: từ vựng tiếng Việt (word-level)
            - 30% trọng số: ký tự có dấu (char-level)
            - Xử lý viết tắt: AI, IT, etc. trong context tiếng Việt
        """
        if not text:
            return 0.0

        # 1. WORD-LEVEL ANALYSIS (70% trọng số)
        words = re.findall(r"\w+", text.lower())
        vietnamese_words_set = self._get_vietnamese_words()

        if not words:
            return 0.0

        # Xử lý các từ viết tắt thông dụng trong tiếng Việt
        vietnamese_abbreviations = self._get_vietnamese_abbreviations()

        vietnamese_word_count = 0
        for word in words:
            if word in vietnamese_words_set:
                vietnamese_word_count += 1
            elif (
                word in vietnamese_abbreviations and len(words) > 2
            ):  # Chỉ tính nếu có context
                vietnamese_word_count += 0.5  # Tính 50% cho viết tắt

        word_score = vietnamese_word_count / len(words)

        # 2. CHAR-LEVEL ANALYSIS (30% trọng số)
        vietnamese_chars = (
            r"[àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđĐ]"
        )
        vietnamese_char_count = len(re.findall(vietnamese_chars, text, re.IGNORECASE))
        total_chars = len(
            re.findall(
                r"[a-zA-ZàáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđĐ]",
                text,
            )
        )

        char_score = vietnamese_char_count / total_chars if total_chars > 0 else 0.0

        # 3. COMBINED SCORE với trọng số
        combined_score = (word_score * 0.7) + (char_score * 0.3)

        # Giới hạn trong khoảng [0, 1]
        return min(1.0, combined_score)

    def clean_vietnamese_text(self, text: str) -> str:
        """
        Làm sạch và chuẩn hóa text tiếng Việt

        Args:
            text (str): Text gốc cần làm sạch

        Returns:
            str: Text đã được làm sạch và chuẩn hóa

        Các bước xử lý:
            1. Chuẩn hóa khoảng trắng (loại bỏ space/tab thừa)
            2. Giảm xuống dòng liên tiếp (từ 3+ xuống 2)
            3. Giữ lại ký tự tiếng Việt và dấu câu cơ bản
            4. Sửa lỗi dính câu (thêm space sau dấu câu)
            5. Loại bỏ space thừa cuối cùng

        Examples:
            >>> parser.clean_vietnamese_text("  Xin   chào.Tôi là AI  ")
            "Xin chào. Tôi là AI"
        """
        if not text:
            return ""

        # Chuẩn hoá space và tab
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Giữ ký tự tiếng Việt, số, chữ cái và dấu câu cơ bản
        text = re.sub(r"[^\w\s\u00C0-\u024F\u1E00-\u1EFF.,!?;:()\[\]\"'\n-]", " ", text)

        # Fix lỗi dính câu (thêm space sau dấu câu nếu thiếu)
        text = re.sub(
            r"([.!?])\s*([A-ZÁÀẮÂÉÈÊÔƠÚÙƯĐÍÌ])",
            r"\1 \2",
            text,
            flags=re.IGNORECASE,
        )

        # Loại bỏ space thừa
        return re.sub(r" +", " ", text).strip()

    def detect_language(self, text: str) -> str:
        """
        Phát hiện ngôn ngữ của text

        Args:
            text (str): Text cần phân tích

        Returns:
            str: Mã ngôn ngữ:
                - "vi": Tiếng Việt (có dấu tiếng Việt)
                - "en": Tiếng Anh (không có dấu, chỉ ký tự ASCII)
                - "mixed": Hỗn hợp
                - "unknown": Không xác định

        Examples:
            >>> parser.detect_language("Xin chào")
            "vi"
            >>> parser.detect_language("Hello world")
            "en"
        """
        if not text or not text.strip():
            return "unknown"

        vn_score = self._calculate_vietnamese_score(text)
        contains_vn = self._contains_vietnamese(text)

        # Đếm ký tự ASCII thuần (a-z, A-Z)
        ascii_chars = len(re.findall(r"[a-zA-Z]", text))
        total_chars = len(
            re.findall(
                r"[a-zA-ZàáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđĐ]",
                text,
            )
        )

        if total_chars == 0:
            return "unknown"

        ascii_ratio = ascii_chars / total_chars

        # Logic phát hiện ngôn ngữ

        # Kiểm tra từ tiếng Anh rõ ràng
        english_words = len(
            re.findall(
                r"\b(speak|english|hello|world|the|and|is|are|this|that|with|for)\b",
                text.lower(),
            )
        )
        total_words = len(text.split())
        english_word_ratio = english_words / total_words if total_words > 0 else 0

        if (
            contains_vn and english_word_ratio > 0.2
        ):  # Có cả tiếng Việt và từ tiếng Anh rõ ràng
            return "mixed"
        elif contains_vn and vn_score > 0.1:  # Có dấu tiếng Việt rõ ràng
            return "vi"
        elif ascii_ratio > 0.9 and not contains_vn:  # Chủ yếu ASCII, không có dấu
            return "en"
        else:
            return "unknown"

    def detect_content_type(self, text: str) -> str:
        """
        Phát hiện loại nội dung của đoạn text

        Args:
            text (str): Đoạn text cần phân loại

        Returns:
            str: Loại nội dung, một trong các giá trị:
                - "table": Nội dung dạng bảng
                - "image": Tham chiếu đến hình ảnh/biểu đồ
                - "header": Tiêu đề/đầu mục
                - "list": Danh sách có đánh số hoặc bullet
                - "text": Text bình thường

        Examples:
            >>> parser.detect_content_type("CHƯƠNG 1: GIỚI THIỆU")
            "header"
            >>> parser.detect_content_type("1. Mục đầu\\n2. Mục hai")
            "list"
        """
        if self._looks_like_table(text):
            return "table"
        if self._looks_like_image_content(text):
            return "image"
        if self._looks_like_header(text):
            return "header"
        if self._looks_like_list(text):
            return "list"
        return "text"

    def _looks_like_table(self, text: str) -> bool:
        """
        Kiểm tra text có giống nội dung bảng không

        Args:
            text (str): Text cần kiểm tra

        Returns:
            bool: True nếu có vẻ là bảng, False nếu không

        Logic: Nếu >= 50% dòng có separator (space nhiều, tab, |) thì coi là bảng
        """
        lines = text.split("\n")
        if len(lines) < 2:
            return False
        sep = sum(1 for l in lines if re.search(r"\s{2,}|\t|[|]", l))
        return sep >= len(lines) * 0.5

    def _looks_like_image_content(self, text: str) -> bool:
        """
        Kiểm tra text có phải tham chiếu hình ảnh/biểu đồ không

        Args:
            text (str): Text cần kiểm tra

        Returns:
            bool: True nếu có từ khóa liên quan đến hình ảnh
        """
        image_keywords = [
            "hình",
            "ảnh",
            "biểu đồ",
            "chart",
            "figure",
            "graph",
            "diagram",
        ]
        return any(word in text.lower() for word in image_keywords)

    def _looks_like_header(self, text: str) -> bool:
        """
        Kiểm tra text có phải tiêu đề không

        Args:
            text (str): Text cần kiểm tra

        Returns:
            bool: True nếu có vẻ là tiêu đề

        Logic:
        - Text ngắn (≤15 từ) và toàn chữ hoa hoặc Title Case
        - KHÔNG phải list (không có bullet points hoặc numbering)
        """
        # Nếu có pattern của list, không phải header
        if re.search(r"(^|\n)\s*(-|\d+\.|[a-zA-Z]\.)", text):
            return False

        # Header phải ngắn và có format đặc biệt
        if len(text.split()) <= 15 and (text.isupper() or text.istitle()):
            return True
        return False

    def _looks_like_list(self, text: str) -> bool:
        """
        Kiểm tra text có phải danh sách không

        Args:
            text (str): Text cần kiểm tra

        Returns:
            bool: True nếu có vẻ là danh sách

        Logic: >= 50% dòng bắt đầu bằng bullet (-), số (1.), hoặc chữ (a.)
        """
        lines = text.split("\n")
        list_lines = sum(
            1 for l in lines if re.match(r"^(-|\d+\.|[a-zA-Z]\.)", l.strip())
        )
        return list_lines >= len(lines) * 0.5

    def extract_metadata(self, text: str, meta: Dict[str, Any]) -> Dict[str, Any]:
        """
        Tạo metadata kèm thông tin cơ bản về text

        Args:
            text (str): Nội dung text
            meta (Dict[str, Any]): Metadata gốc từ loader

        Returns:
            Dict[str, Any]: Metadata được bổ sung với:
                - length: Số ký tự
                - words: Số từ
                - content_type: Loại nội dung
                - language: Ngôn ngữ được phát hiện (vi/en/mixed/unknown)
                - detected: Ngôn ngữ đã được phát hiện (language code)
                - score: Tỷ lệ tiếng Việt (0.0-1.0)
                - plus các metadata gốc
        """
        language = self.detect_language(text)
        return {
            **meta,
            "length": len(text),
            "words": len(text.split()),
            "content_type": self.detect_content_type(text),
            "language": language,
            "detected": language,
            "score": self._calculate_vietnamese_score(text),
        }

    def load_pdf(self, path: str) -> List[Document]:
        """
        Load file PDF với strategy đa tầng

        Args:
            path (str): Đường dẫn đến file PDF

        Returns:
            List[Document]: Danh sách các Document đã extract,
                           mỗi Document chứa text từ 1 page hoặc 1 bảng

        Strategy:
            1. Ưu tiên: PyMuPDF (extract text + layout)
            2. Fallback: PyPDFLoader (extract text cơ bản)
            3. Bổ sung: pdfplumber (extract bảng riêng)

        Raises:
            Log error và return empty list
        """
        docs = []

        # Thử PyMuPDF trước (tốt nhất)
        if PYMUPDF_AVAILABLE:
            docs = self._extract_pymupdf(path)

        # Nếu không có PyMuPDF hoặc thất bại, dùng PyPDFLoader
        if not docs:
            try:
                loader = PyPDFLoader(path)
                for d in loader.load():
                    cleaned = self.clean_vietnamese_text(d.page_content)
                    meta = self.extract_metadata(cleaned, d.metadata)
                    docs.append(Document(page_content=cleaned, metadata=meta))
            except Exception as e:
                logger.error(f"Lỗi PyPDFLoader với {path}: {e}")

        # Thêm bảng từ pdfplumber (nếu có)
        if PDFPLUMBER_AVAILABLE:
            table_docs = self._extract_tables(path)
            docs.extend(table_docs)

        return docs

    def _extract_pymupdf(self, path: str) -> List[Document]:
        """
        Dùng PyMuPDF để extract text từ PDF (method nội bộ)

        Args:
            path (str): Đường dẫn file PDF

        Returns:
            List[Document]: List các Document, mỗi cái = 1 page

        Note: Dùng API của PyMuPDF
        """
        docs = []
        try:
            import fitz

            doc = fitz.open(path)
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                # Extract text từ page
                textpage = page.get_textpage()
                text = textpage.extractText()
                if text.strip():
                    cleaned = self.clean_vietnamese_text(text)
                    meta = self.extract_metadata(
                        cleaned,
                        {
                            "page": page_num,
                            "source": path,
                            "extraction_method": "pymupdf",
                        },
                    )
                    docs.append(Document(page_content=cleaned, metadata=meta))
            doc.close()
        except Exception as e:
            logger.error(f"Lỗi PyMuPDF với {path}: {e}")
        return docs

    def _extract_tables(self, path: str) -> List[Document]:
        """
        Dùng pdfplumber để extract bảng từ PDF (method nội bộ)

        Args:
            path (str): Đường dẫn file PDF

        Returns:
            List[Document]: List các Document, mỗi cái = 1 bảng

        Note: Chỉ extract những bảng có thực sự có nội dung
        """
        docs = []
        try:
            with pdfplumber.open(path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    tables = page.extract_tables()
                    if tables:
                        for table_num, tb in enumerate(tables):
                            rows = ["\t".join([c or "" for c in r]) for r in tb]
                            text = "\n".join(rows)
                            cleaned = self.clean_vietnamese_text(text)
                            meta = self.extract_metadata(
                                cleaned,
                                {
                                    "type": "table",
                                    "page": page_num,
                                    "table_num": table_num,
                                    "source": path,
                                    "extraction_method": "pdfplumber",
                                },
                            )
                            docs.append(Document(page_content=cleaned, metadata=meta))
        except Exception as e:
            logger.error(f"Lỗi pdfplumber với {path}: {e}")
        return docs

    def load_docx(self, path: str) -> List[Document]:
        """
        Load file DOCX (Microsoft Word)

        Args:
            path (str): Đường dẫn đến file .docx

        Returns:
            List[Document]: Danh sách Documents:
                - Mỗi paragraph không rỗng = 1 Document
                - Mỗi table = 1 Document riêng

        Metadata bao gồm:
            - paragraph_num/table_num: Số thứ tự
            - style: Tên style của paragraph (Normal, Heading 1, etc.)
            - type: "paragraph" hoặc "table"
            - source: Đường dẫn file gốc
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx không có, không thể xử lý DOCX")
            return []

        docs = []
        try:
            doc = DocxDocument(path)

            # Extract paragraphs
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    cleaned = self.clean_vietnamese_text(paragraph.text)
                    meta = self.extract_metadata(
                        cleaned,
                        {
                            "paragraph_num": para_num,
                            "style": (
                                paragraph.style.name if paragraph.style else "Normal"
                            ),
                            "source": path,
                            "type": "paragraph",
                            "extraction_method": "python-docx",
                        },
                    )
                    docs.append(Document(page_content=cleaned, metadata=meta))

            # Extract tables
            for table_num, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    table_text.append(row_text)

                if table_text:
                    text = "\n".join(table_text)
                    cleaned = self.clean_vietnamese_text(text)
                    meta = self.extract_metadata(
                        cleaned,
                        {
                            "table_num": table_num,
                            "source": path,
                            "type": "table",
                            "extraction_method": "python-docx",
                        },
                    )
                    docs.append(Document(page_content=cleaned, metadata=meta))

        except Exception as e:
            logger.error(f"Lỗi khi xử lý DOCX {path}: {e}")

        return docs

    def load_txt(self, path: str) -> List[Document]:
        """
        Load file text (.txt, .md) với auto-detect encoding

        Args:
            path (str): Đường dẫn đến file text

        Returns:
            List[Document]: Thường chỉ 1 Document chứa toàn bộ nội dung file

        Encoding fallback:
            1. UTF-8 (mặc định)
            2. UTF-8 with BOM
            3. CP1252 (Windows Latin-1)
            4. ISO-8859-1 (fallback cuối)

        Note: Metadata sẽ chứa encoding đã sử dụng thành công
        """
        docs = []
        encodings_to_try = ["utf-8", "utf-8-sig", "cp1252", "iso-8859-1"]

        for encoding in encodings_to_try:
            try:
                loader = TextLoader(path, encoding=encoding)
                for d in loader.load():
                    cleaned = self.clean_vietnamese_text(d.page_content)
                    meta = self.extract_metadata(
                        cleaned,
                        {
                            **d.metadata,
                            "encoding": encoding,
                            "extraction_method": "langchain-textloader",
                        },
                    )
                    docs.append(Document(page_content=cleaned, metadata=meta))
                break  # Success, exit loop
            except UnicodeDecodeError:
                continue  # Thử encoding tiếp theo
            except Exception as e:
                logger.error(f"Lỗi khi load {path} với encoding {encoding}: {e}")
                continue

        if not docs:
            logger.error(f"Không thể load file {path} với bất kỳ encoding nào")

        return docs

    def load_csv(self, path: str) -> List[Document]:
        """
        Load file CSV

        Args:
            path (str): Đường dẫn đến file .csv

        Returns:
            List[Document]: 1 Document chứa dữ liệu CSV dạng bảng

        Format output: Mỗi row thành 1 dòng, columns cách nhau bằng tab
        """
        docs = []
        encodings_to_try = ["utf-8", "utf-8-sig", "cp1252", "iso-8859-1"]

        for encoding in encodings_to_try:
            try:
                with open(path, "r", encoding=encoding, newline="") as csvfile:
                    # Auto-detect delimiter
                    sample = csvfile.read(1024)
                    csvfile.seek(0)
                    sniffer = csv.Sniffer()
                    delimiter = sniffer.sniff(sample).delimiter

                    reader = csv.reader(csvfile, delimiter=delimiter)
                    rows = []
                    for row in reader:
                        rows.append("\t".join(row))

                    if rows:
                        text = "\n".join(rows)
                        cleaned = self.clean_vietnamese_text(text)
                        meta = self.extract_metadata(
                            cleaned,
                            {
                                "source": path,
                                "type": "csv_table",
                                "encoding": encoding,
                                "delimiter": delimiter,
                                "rows": len(rows),
                                "extraction_method": "csv-reader",
                            },
                        )
                        docs.append(Document(page_content=cleaned, metadata=meta))
                break
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Lỗi CSV {path} với encoding {encoding}: {e}")
                continue

        if not docs:
            logger.error(f"Không thể load CSV {path}")

        return docs

    def load_html(self, path: str) -> List[Document]:
        """
        Load file HTML cơ bản

        Args:
            path (str): Đường dẫn đến file .html/.htm

        Returns:
            List[Document]: 1 Document chứa text đã extract từ HTML

        Note: Chỉ extract text, bỏ qua tags và scripts
        """
        if not BS4_AVAILABLE:
            logger.error("BeautifulSoup không có, không thể xử lý HTML")
            return []

        docs = []
        encodings_to_try = ["utf-8", "utf-8-sig", "cp1252", "iso-8859-1"]

        for encoding in encodings_to_try:
            try:
                with open(path, "r", encoding=encoding) as f:
                    content = f.read()

                soup = BeautifulSoup(content, "html.parser")

                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()

                # Extract text
                text = soup.get_text()

                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (
                    phrase.strip() for line in lines for phrase in line.split("  ")
                )
                text = " ".join(chunk for chunk in chunks if chunk)

                if text.strip():
                    cleaned = self.clean_vietnamese_text(text)
                    meta = self.extract_metadata(
                        cleaned,
                        {
                            "source": path,
                            "type": "html",
                            "encoding": encoding,
                            "extraction_method": "beautifulsoup",
                        },
                    )
                    docs.append(Document(page_content=cleaned, metadata=meta))
                break
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Lỗi HTML {path} với encoding {encoding}: {e}")
                continue

        if not docs:
            logger.error(f"Không thể load HTML {path}")

        return docs

    def load_single_document(self, path: str) -> List[Document]:
        """
        Load một file duy nhất (auto-detect format)

        Args:
            path (str): Đường dẫn đến file cần load

        Returns:
            List[Document]: Danh sách Documents được extract từ file
                           Trả về empty list nếu không thể xử lý

        Supported formats:
            - .pdf: PDF files
            - .docx: Microsoft Word
            - .txt, .md: Text files
            - .csv: CSV files
            - .html, .htm: HTML files (nếu có BeautifulSoup)

        Note: Method này tự động phát hiện định dạng dựa trên extension
        """
        if not os.path.exists(path):
            logger.error(f"Không tìm thấy file: {path}")
            return []

        ext = Path(path).suffix.lower()

        if ext == ".pdf":
            return self.load_pdf(path)
        elif ext in [".txt", ".md"]:
            return self.load_txt(path)
        elif ext == ".docx":
            return self.load_docx(path)
        elif ext == ".csv":
            return self.load_csv(path)
        elif ext in [".html", ".htm"]:
            return self.load_html(path)
        else:
            logger.error(f"Không hỗ trợ định dạng: {ext}")
            logger.info(
                f"Định dạng được hỗ trợ: {', '.join(sorted(self.supported_formats))}"
            )
            return []

    def load_documents(self, files: List[str]) -> List[Document]:
        """
        Load nhiều file cùng lúc

        Args:
            files (List[str]): Danh sách đường dẫn các file cần load

        Returns:
            List[Document]: Danh sách tất cả Documents từ tất cả files
                           Documents từ file nào thất bại sẽ bị bỏ qua

        Note:
            - Xử lý tuần tự từng file
            - File nào lỗi sẽ bị skip, không ảnh hưởng file khác
            - Log thống kê cuối cùng về số documents và files

        Examples:
            >>> docs = parser.load_documents(["file1.pdf", "file2.docx"])
            >>> len(docs)  # Tổng số documents từ cả 2 files
        """
        all_docs = []
        successful_files = 0

        for f in files:
            try:
                docs = self.load_single_document(f)
                if docs:
                    all_docs.extend(docs)
                    successful_files += 1
            except Exception as e:
                logger.error(f"Lỗi khi xử lý file {f}: {e}")

        logger.info(
            f"Đã load {len(all_docs)} documents từ {successful_files}/{len(files)} files"
        )
        return all_docs
